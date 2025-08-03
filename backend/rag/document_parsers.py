"""
Document parsers for various file formats
"""
import os
import mimetypes
import chardet
from typing import Dict, Any, List, Optional, Tuple
from abc import ABC, abstractmethod
import logging
from io import BytesIO

# Import libraries conditionally
try:
    import pypdf
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("pypdf not available. PDF parsing will be disabled.")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX parsing will be disabled.")

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    logging.warning("beautifulsoup4 not available. HTML parsing will be limited.")

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    logging.warning("markdown not available. Markdown parsing will be limited.")

logger = logging.getLogger(__name__)


class ParsedDocument:
    """Container for parsed document data"""
    
    def __init__(self, 
                 content: str,
                 metadata: Dict[str, Any],
                 sections: Optional[List[Dict[str, Any]]] = None):
        self.content = content
        self.metadata = metadata
        self.sections = sections or []
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'content': self.content,
            'metadata': self.metadata,
            'sections': self.sections
        }


class DocumentParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse document and return structured content"""
        pass
    
    @abstractmethod
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if parser can handle the file"""
        pass


class TextParser(DocumentParser):
    """Parser for plain text files"""
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse text file"""
        if file_path:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
        else:
            raw_content = file_content
        
        # Detect encoding
        detected = chardet.detect(raw_content)
        encoding = detected['encoding'] or 'utf-8'
        
        try:
            content = raw_content.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to utf-8 with error handling
            content = raw_content.decode('utf-8', errors='ignore')
        
        metadata = {
            'type': 'text',
            'encoding': encoding,
            'size': len(raw_content),
            'file_path': file_path
        }
        
        return ParsedDocument(content, metadata)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is text"""
        if mime_type:
            return mime_type.startswith('text/')
        
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            return ext in ['.txt', '.text', '.log', '.csv', '.tsv']
        
        return False


class PDFParser(DocumentParser):
    """Parser for PDF files"""
    
    def __init__(self):
        if not PDF_AVAILABLE:
            raise ImportError("pypdf is required for PDF parsing. Install with: pip install pypdf")
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse PDF file"""
        if file_path:
            reader = pypdf.PdfReader(file_path)
        else:
            reader = pypdf.PdfReader(BytesIO(file_content))
        
        # Extract text from all pages
        full_text = []
        sections = []
        
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text.strip():
                full_text.append(page_text)
                sections.append({
                    'type': 'page',
                    'page_number': page_num + 1,
                    'content': page_text,
                    'metadata': {
                        'page': page_num + 1,
                        'total_pages': len(reader.pages)
                    }
                })
        
        # Extract metadata
        metadata = {
            'type': 'pdf',
            'pages': len(reader.pages),
            'file_path': file_path
        }
        
        # Try to extract PDF metadata
        if reader.metadata:
            pdf_meta = {}
            if reader.metadata.title:
                pdf_meta['title'] = reader.metadata.title
            if reader.metadata.author:
                pdf_meta['author'] = reader.metadata.author
            if reader.metadata.subject:
                pdf_meta['subject'] = reader.metadata.subject
            if reader.metadata.creator:
                pdf_meta['creator'] = reader.metadata.creator
            if reader.metadata.creation_date:
                pdf_meta['creation_date'] = str(reader.metadata.creation_date)
            
            metadata['pdf_metadata'] = pdf_meta
        
        content = '\n\n'.join(full_text)
        return ParsedDocument(content, metadata, sections)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is PDF"""
        if mime_type:
            return mime_type == 'application/pdf'
        
        if file_path:
            return file_path.lower().endswith('.pdf')
        
        return False


class DOCXParser(DocumentParser):
    """Parser for DOCX files"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse DOCX file"""
        if file_path:
            doc = docx.Document(file_path)
        else:
            doc = docx.Document(BytesIO(file_content))
        
        # Extract text and structure
        full_text = []
        sections = []
        current_section = None
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = element
                text = ''
                for run in para.iter():
                    if run.tag.endswith('t'):
                        text += run.text
                
                if text.strip():
                    # Check if it's a heading
                    style = None
                    for style_elem in para.iter():
                        if style_elem.tag.endswith('pStyle'):
                            style = style_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            break
                    
                    if style and 'heading' in style.lower():
                        # Save previous section
                        if current_section and current_section['content']:
                            sections.append(current_section)
                        
                        # Start new section
                        current_section = {
                            'type': 'section',
                            'title': text.strip(),
                            'level': int(style[-1]) if style[-1].isdigit() else 1,
                            'content': [],
                            'metadata': {'style': style}
                        }
                    else:
                        # Regular paragraph
                        full_text.append(text)
                        if current_section:
                            current_section['content'].append(text)
            
            elif element.tag.endswith('tbl'):  # Table
                # Extract table text
                table_text = []
                for row in element.iter():
                    if row.tag.endswith('tr'):
                        row_text = []
                        for cell in row.iter():
                            if cell.tag.endswith('tc'):
                                cell_text = ''
                                for t in cell.iter():
                                    if t.tag.endswith('t'):
                                        cell_text += t.text
                                if cell_text:
                                    row_text.append(cell_text.strip())
                        if row_text:
                            table_text.append(' | '.join(row_text))
                
                if table_text:
                    table_content = '\n'.join(table_text)
                    full_text.append(table_content)
                    if current_section:
                        current_section['content'].append(table_content)
        
        # Add last section
        if current_section and current_section['content']:
            sections.append(current_section)
        
        # Convert sections content to text
        for section in sections:
            if isinstance(section['content'], list):
                section['content'] = '\n\n'.join(section['content'])
        
        metadata = {
            'type': 'docx',
            'sections_count': len(sections),
            'file_path': file_path
        }
        
        # Extract document properties
        if doc.core_properties:
            doc_props = {}
            if doc.core_properties.title:
                doc_props['title'] = doc.core_properties.title
            if doc.core_properties.author:
                doc_props['author'] = doc.core_properties.author
            if doc.core_properties.subject:
                doc_props['subject'] = doc.core_properties.subject
            if doc.core_properties.created:
                doc_props['created'] = str(doc.core_properties.created)
            if doc.core_properties.modified:
                doc_props['modified'] = str(doc.core_properties.modified)
            
            metadata['document_properties'] = doc_props
        
        content = '\n\n'.join(full_text)
        return ParsedDocument(content, metadata, sections)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is DOCX"""
        if mime_type:
            return mime_type in [
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'application/msword'
            ]
        
        if file_path:
            return file_path.lower().endswith(('.docx', '.doc'))
        
        return False


class HTMLParser(DocumentParser):
    """Parser for HTML files"""
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse HTML file"""
        if file_path:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
        else:
            raw_content = file_content
        
        # Detect encoding
        detected = chardet.detect(raw_content)
        encoding = detected['encoding'] or 'utf-8'
        
        try:
            html_content = raw_content.decode(encoding)
        except UnicodeDecodeError:
            html_content = raw_content.decode('utf-8', errors='ignore')
        
        if BS4_AVAILABLE:
            # Use BeautifulSoup for better parsing
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Extract sections based on headers
            sections = []
            headers = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
            
            for i, header in enumerate(headers):
                # Find content between this header and the next
                content_elements = []
                current = header.next_sibling
                
                while current and (i + 1 >= len(headers) or current != headers[i + 1]):
                    if hasattr(current, 'get_text'):
                        content_elements.append(current.get_text(strip=True))
                    current = current.next_sibling if hasattr(current, 'next_sibling') else None
                
                sections.append({
                    'type': 'section',
                    'title': header.get_text(strip=True),
                    'level': int(header.name[1]),
                    'content': '\n'.join(filter(None, content_elements)),
                    'metadata': {'tag': header.name}
                })
            
            # Extract metadata
            metadata = {
                'type': 'html',
                'encoding': encoding,
                'file_path': file_path
            }
            
            # Extract meta tags
            meta_tags = {}
            for meta in soup.find_all('meta'):
                if meta.get('name'):
                    meta_tags[meta['name']] = meta.get('content', '')
                elif meta.get('property'):
                    meta_tags[meta['property']] = meta.get('content', '')
            
            if meta_tags:
                metadata['meta_tags'] = meta_tags
            
            # Extract title
            if soup.title:
                metadata['title'] = soup.title.string
            
        else:
            # Basic regex-based extraction
            import re
            
            # Remove script and style
            text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', ' ', text)
            
            # Clean up whitespace
            text = ' '.join(text.split())
            
            sections = []
            metadata = {
                'type': 'html',
                'encoding': encoding,
                'file_path': file_path,
                'parser': 'regex'
            }
        
        return ParsedDocument(text, metadata, sections)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is HTML"""
        if mime_type:
            return mime_type in ['text/html', 'application/xhtml+xml']
        
        if file_path:
            return file_path.lower().endswith(('.html', '.htm', '.xhtml'))
        
        return False


class MarkdownParser(DocumentParser):
    """Parser for Markdown files"""
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse Markdown file"""
        if file_path:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
        else:
            raw_content = file_content
        
        # Decode content
        try:
            md_content = raw_content.decode('utf-8')
        except UnicodeDecodeError:
            md_content = raw_content.decode('utf-8', errors='ignore')
        
        # Extract sections based on headers
        import re
        sections = []
        lines = md_content.split('\n')
        current_section = None
        
        for line in lines:
            # Check if line is a header
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            
            if header_match:
                # Save previous section
                if current_section:
                    current_section['content'] = '\n'.join(current_section['content']).strip()
                    sections.append(current_section)
                
                # Start new section
                level = len(header_match.group(1))
                title = header_match.group(2).strip()
                current_section = {
                    'type': 'section',
                    'title': title,
                    'level': level,
                    'content': [],
                    'metadata': {'markdown_level': level}
                }
            elif current_section:
                current_section['content'].append(line)
        
        # Add last section
        if current_section:
            current_section['content'] = '\n'.join(current_section['content']).strip()
            sections.append(current_section)
        
        # Convert to HTML if markdown library is available
        if MARKDOWN_AVAILABLE:
            html_content = markdown.markdown(md_content, extensions=['extra', 'codehilite'])
            # Extract plain text from HTML
            if BS4_AVAILABLE:
                soup = BeautifulSoup(html_content, 'html.parser')
                text = soup.get_text()
            else:
                # Basic HTML tag removal
                text = re.sub(r'<[^>]+>', '', html_content)
        else:
            # Use original markdown text
            text = md_content
        
        metadata = {
            'type': 'markdown',
            'sections_count': len(sections),
            'file_path': file_path
        }
        
        return ParsedDocument(text, metadata, sections)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is Markdown"""
        if mime_type:
            return mime_type in ['text/markdown', 'text/x-markdown']
        
        if file_path:
            return file_path.lower().endswith(('.md', '.markdown', '.mdown'))
        
        return False


class CodeParser(DocumentParser):
    """Parser for source code files"""
    
    LANGUAGE_EXTENSIONS = {
        '.py': 'python',
        '.js': 'javascript',
        '.ts': 'typescript',
        '.java': 'java',
        '.cpp': 'cpp',
        '.c': 'c',
        '.cs': 'csharp',
        '.php': 'php',
        '.rb': 'ruby',
        '.go': 'go',
        '.rs': 'rust',
        '.swift': 'swift',
        '.kt': 'kotlin',
        '.scala': 'scala',
        '.r': 'r',
        '.m': 'matlab',
        '.sql': 'sql',
        '.sh': 'bash',
        '.ps1': 'powershell',
        '.yaml': 'yaml',
        '.yml': 'yaml',
        '.json': 'json',
        '.xml': 'xml',
        '.html': 'html',
        '.css': 'css',
        '.scss': 'scss',
        '.vue': 'vue',
        '.jsx': 'jsx',
        '.tsx': 'tsx'
    }
    
    def parse(self, file_path: str = None, file_content: bytes = None) -> ParsedDocument:
        """Parse source code file"""
        if file_path:
            with open(file_path, 'rb') as f:
                raw_content = f.read()
            
            # Detect language from extension
            ext = os.path.splitext(file_path)[1].lower()
            language = self.LANGUAGE_EXTENSIONS.get(ext, 'unknown')
        else:
            raw_content = file_content
            language = 'unknown'
        
        # Decode content
        detected = chardet.detect(raw_content)
        encoding = detected['encoding'] or 'utf-8'
        
        try:
            code_content = raw_content.decode(encoding)
        except UnicodeDecodeError:
            code_content = raw_content.decode('utf-8', errors='ignore')
        
        metadata = {
            'type': 'code',
            'language': language,
            'encoding': encoding,
            'size': len(raw_content),
            'lines': len(code_content.splitlines()),
            'file_path': file_path
        }
        
        # For code files, we don't extract sections as the structure
        # should be preserved for the code chunker
        return ParsedDocument(code_content, metadata)
    
    def can_parse(self, file_path: str = None, mime_type: str = None) -> bool:
        """Check if file is source code"""
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            return ext in self.LANGUAGE_EXTENSIONS
        
        if mime_type:
            code_mimes = [
                'text/x-python', 'text/x-java', 'text/x-c',
                'application/javascript', 'application/json',
                'application/xml', 'text/x-sql'
            ]
            return any(mime in mime_type for mime in code_mimes)
        
        return False


class DocumentParserFactory:
    """Factory for creating appropriate document parsers"""
    
    _parsers = [
        PDFParser,
        DOCXParser,
        HTMLParser,
        MarkdownParser,
        CodeParser,
        TextParser  # TextParser should be last as fallback
    ]
    
    @classmethod
    def get_parser(cls, file_path: str = None, mime_type: str = None) -> Optional[DocumentParser]:
        """Get appropriate parser for file"""
        for parser_class in cls._parsers:
            try:
                parser = parser_class()
                if parser.can_parse(file_path, mime_type):
                    return parser
            except ImportError:
                # Parser not available due to missing dependencies
                continue
        
        # Default to text parser
        return TextParser()
    
    @classmethod
    def parse_document(cls, file_path: str = None, 
                      file_content: bytes = None,
                      mime_type: str = None) -> ParsedDocument:
        """Parse document using appropriate parser"""
        if file_path and not mime_type:
            mime_type, _ = mimetypes.guess_type(file_path)
        
        parser = cls.get_parser(file_path, mime_type)
        if parser:
            return parser.parse(file_path, file_content)
        
        raise ValueError(f"No parser available for file: {file_path or 'content'}")


# Convenience function
def parse_document(file_path: str = None, 
                  file_content: bytes = None,
                  mime_type: str = None) -> ParsedDocument:
    """Parse document using appropriate parser"""
    return DocumentParserFactory.parse_document(file_path, file_content, mime_type)